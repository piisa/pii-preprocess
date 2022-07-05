"""
Read Microsoft Word documents (docx)
"""

from datetime import datetime

from docx import Document
from docx.text.paragraph import Paragraph

from typing import Dict, Iterable, Tuple

from pii_data.types.document import TreeSrcDocument, SequenceSrcDocument, TYPE_META
from pii_data.types.localdoc import dump_file


def add_subchunk(parent: Dict, subchunk: Dict):
    """
    Add a child subchunk to a chunk
    """
    if parent is None:
        parent = {}
    if 'chunks' not in parent:
        parent['chunks'] = []
    parent['chunks'].append(subchunk)


def iter_paragraphs(it: Iterable[Paragraph]) -> Iterable[Tuple[str, str]]:
    """
    Return paragraphs from the document. Will merge paragraphs containing
    only whitespace with the previous one.
     :param it: an iterable over the Word paragraphs
     :return: a tuple (paragraph-text, paragraph-style)
    """
    prev = None
    for p in it:
        cur = p.text + "\n", p.style.name
        # If blank, add to previous paragraph and continue
        if not p.text or p.text.isspace():
            prev = (prev[0]+cur[0], prev[1] or cur[1]) if prev else (cur[0], None)
            continue
        # Return the previous paragraph
        if prev:
            yield prev
        # Set the prev paragraph for the next iteration
        prev = cur
    # Return a pending last paragraph
    if prev:
        yield prev


class _TreeReader:
    """
    Iterate over the paragraphs of the document, building a tree by using its
    Heading styles
    """

    def __init__(self, para: Iterable[Paragraph]):
        """
          :param para: an iterable over the document paragraphs
        """
        self.para = para


    def __iter__(self) -> Iterable[Dict]:
        """
        Return an iterable over all 1st level sections; each chunk contains
        the section and all its contents as a subtree.
        Adds "section" context to top-level chunks (the "level" context is
        added automatically by the TreeSrcDocument base class)
        """
        levels = []
        newlevel = 0
        heading = None
        topchunk = newchunk = None
        n = 0

        for text, style in iter_paragraphs(self.para):

            curlevel = newlevel
            heading = style.startswith("Heading")
            if heading:
                try:
                    newlevel = int(style[7:])
                except Exception:
                    heading = False

            n += 1
            newchunk = {"data": text, "id": f"P{n}"}

            # A non-header
            if not heading:
                if curlevel == 0:
                    yield newchunk      # top-level non-heading
                else:
                    add_subchunk(levels[-1], newchunk)
                continue

            # A heading
            if newlevel == 1:                # a level 1 heading

                # The previous top-level group
                if topchunk:
                    yield topchunk

                # Reset top-level group
                topchunk = newchunk
                topchunk["context"] = {"section": text.strip()}
                levels = [topchunk]

            elif newlevel == curlevel:          # heading of the same level

                levels = levels[:curlevel-1]
                add_subchunk(levels[-1], newchunk)
                levels.append(newchunk)

            elif newlevel > curlevel:           # subheading
                if not levels:
                    levels = [newchunk]
                else:
                    for lev in range(curlevel+1, newlevel):
                        add_subchunk(levels[-1], {})
                        levels.append(levels[-1]["chunks"][-1])
                    add_subchunk(levels[-1], newchunk)
                    levels.append(levels[-1]["chunks"][-1])

            elif newlevel < curlevel:           # heading upper in the hierarchy

                levels = levels[:newlevel-1]
                add_subchunk(levels[-1], newchunk)
                levels.append(newchunk)

        # A last top-level group?
        if topchunk:
            yield topchunk


# ------------------------------------------------------------------------


class _BaseMsWordDocument:

    def _open(self, filename: str, doctype: str) -> TYPE_META:
        """
        Open an MS Word file and load it
          :param filename: Word filename to read
          :return: the document general metadata
        """
        # Open the Word file
        self.name = filename
        self.doc = Document(filename)
        #print(self.doc.sections)
        #print(self.doc.settings)

        # Read document generic metadata
        docinfo = {"origin": "msword", "type": doctype}
        for n in ("title", "author", "category", "modified"):
            v = getattr(self.doc.core_properties, n, None)
            if not v:
                continue
            if n != "modified":
                docinfo[n] = v.replace("\n", " ")
            else:
                docinfo["date"] = v.isoformat() if isinstance(v, datetime) else str(v)

        return {"document": docinfo}


    def dump(self, outname: str, **kwargs):
        """
        Dump into a serialized SrcDocument output file
        """
        dump_file(self, outname, **kwargs)



class SeqMsWordDocument(SequenceSrcDocument, _BaseMsWordDocument):
    """
    Read a MS Word document as a Sequence Source Document
    (i.e. flat paragraphs).
    """

    def __init__(self, filename: str, metadata: TYPE_META = None, **kwargs):
        docmeta = self._open(filename, "sequence")
        super().__init__(metadata=docmeta, **kwargs)
        if metadata:
            self.add_metadata(**metadata)


    def iter_base(self) -> Iterable[Dict]:
        for n, (t, _) in enumerate(iter_paragraphs(self.doc.paragraphs), start=1):
            yield {"data": t, "id": f"P{n}"}


class TreeMsWordDocument(TreeSrcDocument, _BaseMsWordDocument):
    """
    Read a MS Word document into a Tree Source Document, by using the Word
    Heading styles to infer the tree.
    """

    def __init__(self, filename: str, metadata: TYPE_META = None, **kwargs):
        docmeta = self._open(filename, "tree")
        super().__init__(metadata=docmeta, **kwargs)
        if metadata:
            self.add_metadata(**metadata)


    def iter_base(self) -> Iterable[Dict]:
        return _TreeReader(self.doc.paragraphs)



class MsWordDocument:
    """
    Wrapper to read a Word file either a a tree or as a flat list
    """

    def __new__(self, filename: str, tree: bool = True, **docinfo):
        """
         :param filename: Word file to read
         :param tree: create a Tree document by using Word headings
         :param docinfo: add to document general metadata
        """
        if tree:
            return TreeMsWordDocument(filename, **docinfo)
        else:
            return SeqMsWordDocument(filename, **docinfo)
